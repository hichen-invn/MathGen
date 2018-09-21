print("====Generate Math Quiz====")

import random
import datetime
from docx import Document
from docx.shared import Pt

now = datetime.datetime.now()

class MathGen(object):
    '''
       For a given params, to generate math addition quizs
       :param type: math quiz type, selection of Adds, Subs, Muls, Divs
       :param quiz_no:  number of total quiz to generate
       :param cal_no:  how many calculations per quiz
       :param dig:  digital number of each calculation number
       :param same_dig:  True to be same digits numbers calculation, False to be one digits less on 2nd and later numbers
       :param: no_neg: True to guarantee no negative number happens, False could see negative number result
       :param q_per_l: quiz to print per line for print
       :param f_size: font size in word file for print
       :param: quizDic: a fictionary to story generated math quiz
       :return:
       '''
    def __init__(self, n):
        self.type = None
        self.quiz_no = n
        self.cal_no = 2
        self.dig = 2
        self.same_dig = False
        self.no_neg = False
        self.q_per_l = 2
        self.f_size = 20
        self.quizDic = {}

    def set_type(self, type):
        self.type = type

    def set_quiz_no(self, quiz_no):
        self.quiz_no = quiz_no

    def set_cal_no(self, cal_no):
        self.cal_no = cal_no

    def set_dig(self, dig):
        self.dig = dig

    def set_same_dig(self, same_dig):
        self.same_dig = same_dig

    def set_no_neg(self, no_neg):
        self.no_neg = no_neg

    def set_q_per_l(self, q_per_l):
        self.q_per_l = q_per_l

    def set_f_size(self, f_size):
        self.f_size = f_size

    def setParams(self, type, quiz_no, cal_no, dig, same_dig, no_neg, q_per_l, f_size):
        """ batch set up parameters """
        self.type = type
        self.quiz_no = quiz_no
        self.cal_no = cal_no
        self.dig = dig
        self.same_dig = same_dig
        self.no_neg = no_neg
        self.q_per_l = q_per_l
        self.f_size = f_size

    def getParams(self):
        """ batch get parameters """
        return self.type, self.quiz_no, self.cal_no, self.dig, self.same_dig, self.no_neg, self.q_per_l, self.f_size

    def mathadds(self):
        self.type = 'Adds'
        self.quizDic = {}
        for i in range(self.quiz_no):
            quiz = []
            upper_bound = 10 ** self.dig - 1
            lower_bound = 10 ** (self.dig-1)
            first_no = random.randint(lower_bound, upper_bound)
            for j in range(self.cal_no):
                if j == 0:
                    quiz.append(str(first_no))
                else:
                    if self.same_dig:
                        quiz.append(str(random.randint(lower_bound, upper_bound)))
                    elif self.dig>1:
                        quiz.append(str(random.randint(1, lower_bound-1)))
                    else:
                        quiz.append(str(random.randint(1, upper_bound)))
                quiz.append(' + ')
            quiz[-1] = ' = '
            self.quizDic['Q'+str(i+1)+':'] = ''.join(quiz)

    def mathsubs(self):
        self.type = 'Subs'
        self.quizDic = {}
        for i in range(self.quiz_no):
            quiz = []
            upper_bound = 10 ** self.dig - 1
            lower_bound = 10 ** (self.dig-1)
            first_no = random.randint(lower_bound, upper_bound)
            for j in range(self.cal_no):
                if j == 0:
                    quiz.append(str(first_no))
                else:
                    if self.same_dig:
                        if self.no_neg:
                            quiz.append(str(random.randint(1, first_no)))
                        else:
                            quiz.append(str(random.randint(lower_bound, upper_bound)))
                    elif self.dig>1:
                        quiz.append(str(random.randint(1, lower_bound-1)))
                    else:
                        quiz.append(str(random.randint(1, first_no)))
                quiz.append(' - ')
            quiz[-1] = ' = '
            self.quizDic['Q'+str(i + 1)+':'] = ''.join(quiz)

    def mathmuls(self):
        self.type = 'Muls'
        self.quizDic = {}
        for i in range(self.quiz_no):
            quiz = []
            upper_bound = 10 ** self.dig - 1
            lower_bound = 10 ** (self.dig - 1)
            first_no = random.randint(lower_bound, upper_bound)
            for j in range(self.cal_no):
                if j == 0:
                    quiz.append(str(first_no))
                else:
                    if self.same_dig:
                        quiz.append(str(random.randint(lower_bound, upper_bound)))
                    elif self.dig > 1:
                        quiz.append(str(random.randint(1, lower_bound - 1)))
                    else:
                        quiz.append(str(random.randint(1, upper_bound)))
                quiz.append(' * ')
            quiz[-1] = ' = '
            self.quizDic['Q' + str(i + 1) + ':'] = ''.join(quiz)

    def savedocs(self):
        fname = 'MathQuizGen_'+now.strftime("%m-%d_")+str(self.cal_no)+'x'+str(self.dig)+'-digit'+self.type
        document = Document()
        document.add_heading(fname, 0)
        style = document.styles['Normal']
        font = style.font
        font.size = Pt(self.f_size)
        style.paragraph_format.space_after = Pt(self.f_size * 1.25)

        table = document.add_table(rows=int(self.quiz_no / self.q_per_l) + 1, cols=self.q_per_l)
        row, col = 0, 0
        for key, item in self.quizDic.items():
            table.rows[row].cells[col].text = key + item
            col += 1
            if col == self.q_per_l:
                col = 0
                row += 1

        document.save(fname + '.docx')
        return


mg = MathGen(0)

mg.set_quiz_no(70)
# mg.set_cal_no(2)
# mg.set_dig(2)
mg.set_same_dig(True)
mg.mathadds()
print('Adds params(type/qn/cn/d/sd/nn/qpl/fs): ', mg.getParams())
mg.savedocs()

mg.set_quiz_no(46)
# mg.set_cal_no(2)
# mg.set_dig(2)
mg.set_same_dig(True)
mg.set_no_neg(True)
mg.mathsubs()
print('Subs params(type/qn/cn/d/sd/nn/qpl/fs): ', mg.getParams())
mg.savedocs()

mg.set_quiz_no(22)
# mg.set_cal_no(2)
mg.set_dig(1)
mg.set_same_dig(True)
mg.mathmuls()
print('Muls params(type/qn/cn/d/sd/nn/qpl/fs): ', mg.getParams())
mg.savedocs()

print("====end====")

